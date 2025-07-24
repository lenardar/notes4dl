# test.py
# 这是一个包含所有功能的单一文件版本

import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
import threading
import requests
import json
from docx import Document
from docx.shared import Pt, RGBColor

# ==============================================================================
# 1. 翻译 API 封装部分
# ==============================================================================

class BaseTranslator:
    """
    翻译服务的基类。
    所有具体的翻译器实现都应继承此类，并实现 translate 方法。
    这种设计使得未来添加新的翻译服务（如 Google Translate, DeepL 等）变得容易。
    """
    def __init__(self, api_key: str):
        if not api_key:
            raise ValueError("API 密钥不能为空。")
        self.api_key = api_key

    def translate(self, text: str, source_lang: str = 'en', target_lang: str = 'zh') -> str:
        """
        翻译给定的文本。
        这是一个抽象方法，子类必须重写此方法以提供具体的翻译逻辑。
        
        :param text: 需要翻译的文本。
        :param source_lang: 源语言代码 (例如, 'en' 代表英语)。
        :param target_lang: 目标语言代码 (例如, 'zh' 代表中文)。
        :return: 翻译后的文本字符串。
        """
        raise NotImplementedError("子类必须实现 translate 方法。")

class SiliconFlowTranslator(BaseTranslator):
    """
    使用“硅基流动” (SiliconFlow) API 的翻译器。
    它通过调用聊天模型并提供特定指令来执行翻译任务。
    """
    def __init__(self, api_key: str):
        super().__init__(api_key)
        # 请根据官方文档确认最新的 API 端点 URL
        self.api_url = "https://api.siliconflow.cn/v1/chat/completions"

    def translate(self, text: str, source_lang: str = 'en', target_lang: str = 'zh') -> str:
        """
        使用硅基流动 API 翻译文本。
        """
        if not text.strip():
            return "" # 如果是空文本或仅含空格，则直接返回空字符串

        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        # 构造符合硅基流动 API 要求的请求体 (payload)。
        # 我们使用 "system" 角色来设定模型的行为，使其成为一个专业的翻译引擎。
        payload = {
            "model": "deepseek-ai/DeepSeek-V3",  # 您可以根据需求更换为其他适合的模型
            "messages": [
                {
                    "role": "system",
                    "content": "你是一个专业的翻译引擎，请将用户提供的英文段落准确地翻译成中文。请直接返回翻译结果，不要包含任何额外的解释、标签或与翻译无关的文字。"
                },
                {
                    "role": "user",
                    "content": text
                }
            ],
            "stream": False,
            "temperature": 0.1 # 较低的温度值使输出更稳定和可预测
        }

        try:
            response = requests.post(self.api_url, headers=headers, data=json.dumps(payload), timeout=30)
            response.raise_for_status()  # 如果 HTTP 状态码表示错误 (4xx 或 5xx)，则抛出异常
            result = response.json()
            
            # 从 API 返回的 JSON 结构中安全地提取翻译内容
            translation = result.get('choices', [{}])[0].get('message', {}).get('content', '')
            return translation.strip()

        except requests.exceptions.RequestException as e:
            error_message = f"翻译 API 请求失败: {e}"
            print(error_message)
            return error_message
        except (KeyError, IndexError, TypeError) as e:
            error_message = f"解析 API 响应失败: {e} - 响应内容: {response.text if 'response' in locals() else '无响应'}"
            print(error_message)
            return error_message

class TranslatorFactory:
    """
    翻译器工厂类。
    用于根据指定的 API 名称创建相应的翻译器实例。
    这种模式将对象的创建与使用分离，提高了代码的灵活性和可维护性。
    """
    @staticmethod
    def get_translator(api_name: str, api_key: str) -> BaseTranslator:
        """
        根据名称获取翻译器实例。
        :param api_name: API 服务的名称 (例如, 'silicon_flow')。
        :param api_key: 该服务所需的 API 密钥。
        :return: 一个 BaseTranslator 的子类实例。
        """
        if api_name.lower() == "silicon_flow":
            return SiliconFlowTranslator(api_key)
        # --- 添加其他 API 的空间 ---
        # elif api_name.lower() == "another_api":
        #     return AnotherApiTranslator(api_key)
        else:
            raise ValueError(f"未知的 API 名称: {api_name}")


# ==============================================================================
# 2. GUI 应用和 Word 处理部分
# ==============================================================================

class TranslatorApp:
    """
    图形用户界面 (GUI) 的主应用程序类。
    负责创建窗口、控件，并处理用户交互。
    """
    def __init__(self, root: tk.Tk):
        self.root = root
        self.root.title("Word 沉浸式翻译工具 (单文件版)")
        self.root.geometry("650x450")

        # --- 配置区 ---
        self.api_service_name = "silicon_flow"
        # 在此处填入您的硅基流动 API Key，或在运行时在界面上输入
        self.default_api_key = "sk-edxhkwriingynarkoxxnqnuydfezydxoyhfzduwchoexwyen"

        # --- 控件创建和布局 ---
        # API 密钥输入框
        tk.Label(root, text="API 密钥 (硅基流动):").pack(pady=(10, 0))
        self.api_key_entry = tk.Entry(root, width=60)
        self.api_key_entry.insert(0, self.default_api_key)
        self.api_key_entry.pack()

        # 文件选择按钮和标签
        tk.Label(root, text="请选择要翻译的 Word 文档 (.docx)").pack(pady=(10, 0))
        self.select_file_button = tk.Button(root, text="选择文件", command=self.select_file)
        self.select_file_button.pack(pady=5)
        self.selected_file_label = tk.Label(root, text="尚未选择文件", fg="blue")
        self.selected_file_label.pack()

        # 开始翻译按钮
        self.translate_button = tk.Button(root, text="开始翻译", command=self.start_translation_thread, state=tk.DISABLED, font=('Helvetica', 10, 'bold'))
        self.translate_button.pack(pady=15)

        # 日志输出区域
        tk.Label(root, text="处理日志:").pack()
        self.log_text = scrolledtext.ScrolledText(root, height=12, state=tk.DISABLED, wrap=tk.WORD, font=("Courier New", 9))
        self.log_text.pack(pady=5, padx=10, fill=tk.BOTH, expand=True)
        
        self.input_file_path = ""

    def select_file(self):
        """处理“选择文件”按钮的点击事件，打开文件对话框。"""
        path = filedialog.askopenfilename(
            title="选择 Word 文档",
            filetypes=(("Word Documents", "*.docx"), ("All files", "*.*"))
        )
        if path:
            self.input_file_path = path
            # 显示文件名，如果文件名太长则截断
            filename = path.split('/')[-1]
            display_text = f"已选择: {filename}" if len(filename) < 60 else f"已选择: ...{filename[-55:]}"
            self.selected_file_label.config(text=display_text)
            self.translate_button.config(state=tk.NORMAL)
            self.log_message(f"准备翻译文件: {self.input_file_path}")

    def log_message(self, msg: str):
        """安全地向日志文本框中添加消息，并自动滚动到底部。"""
        def _log():
            self.log_text.config(state=tk.NORMAL)
            self.log_text.insert(tk.END, msg + "\n")
            self.log_text.see(tk.END)
            self.log_text.config(state=tk.DISABLED)
        # 使用 after 确保 GUI 更新在主线程中执行
        self.root.after(0, _log)

    def set_ui_state(self, enabled: bool):
        """统一设置界面控件的启用/禁用状态。"""
        state = tk.NORMAL if enabled else tk.DISABLED
        self.select_file_button.config(state=state)
        self.translate_button.config(state=state)
        self.api_key_entry.config(state='normal' if enabled else 'readonly')


    def start_translation_thread(self):
        """
        当用户点击“开始翻译”时调用此方法。
        它会验证输入，并启动一个新线程来处理翻译任务，以防止 GUI 冻结。
        """
        api_key = self.api_key_entry.get().strip()
        if not api_key or api_key == "sk-xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx":
            messagebox.showerror("API 密钥无效", "请输入您的有效 API 密钥。")
            return
        
        # 弹出保存文件对话框，让用户选择输出路径
        output_file_path = filedialog.asksaveasfilename(
            title="保存翻译后的文件",
            defaultextension=".docx",
            initialfile=f"[翻译后] {self.input_file_path.split('/')[-1]}",
            filetypes=(("Word Documents", "*.docx"), ("All files", "*.*"))
        )

        if not output_file_path:
            self.log_message("用户取消保存，操作中止。")
            return

        self.set_ui_state(False) # 禁用界面控件
        self.log_message("="*30)
        self.log_message("翻译任务开始，请稍候...")

        # 创建并启动后台线程
        thread = threading.Thread(
            target=self.process_word_document,
            args=(self.input_file_path, output_file_path, api_key)
        )
        thread.daemon = True # 设置为守护线程，主窗口关闭时线程自动退出
        thread.start()

    def process_word_document(self, input_path: str, output_path: str, api_key: str):
        """
        这是在后台线程中运行的核心函数。
        它负责读取、逐段翻译和写入 Word 文档。
        """
        try:
            translator = TranslatorFactory.get_translator(self.api_service_name, api_key)
            document = Document(input_path)
            
            # 过滤掉空的段落，只对有内容的段落进行计数和翻译
            paragraphs_to_translate = [p for p in document.paragraphs if p.text.strip()]
            total_paragraphs = len(paragraphs_to_translate)
            
            self.log_message(f"文档加载成功，共找到 {total_paragraphs} 个非空段落需要翻译。")

            for i, p in enumerate(paragraphs_to_translate):
                original_text = p.text
                self.log_message(f"[{i+1}/{total_paragraphs}] 翻译中: {original_text[:40]}...")
                
                translated_text = translator.translate(original_text)
                
                # 实现沉浸式翻译效果
                # 在原文段落之后添加一个新段落来显示译文
                # 使用 add_paragraph 比 insert_paragraph_before 更直观
                # 但为了保持格式，我们需要找到 p 的下一个元素并插入
                # 一个更稳健的方法是直接修改原段落，然后添加新段落
                
                # 1. 在原段落末尾添加一个换行符，为译文留出空间
                p.add_run() 
                
                # 2. 添加译文段落
                # insert_paragraph_before 可以在循环中安全使用
                translation_p = p.insert_paragraph_before('')
                run = translation_p.add_run(translated_text)
                
                # 3. （可选）为译文设置独特的样式以作区分
                font = run.font
                font.color.rgb = RGBColor(0x00, 0x00, 0x8B) # 深蓝色，比之前的宝蓝色更易读
                font.size = Pt(10.5) # 可根据 Word 默认字体大小调整

            self.log_message("所有段落翻译完毕，正在保存文件...")
            document.save(output_path)
            self.log_message(f"翻译完成！文件已保存至: {output_path}")
            self.log_message("="*30)
            messagebox.showinfo("任务完成", f"文件翻译成功！\n已保存到：{output_path}")

        except Exception as e:
            error_msg = f"处理过程中发生严重错误: {e}"
            self.log_message(f"错误: {error_msg}")
            messagebox.showerror("严重错误", error_msg)
        finally:
            # 确保无论成功还是失败，GUI 控件都会被重新启用
            self.root.after(0, self.set_ui_state, True)

# ==============================================================================
# 3. 程序入口
# ==============================================================================
if __name__ == "__main__":
    # 确保在高DPI屏幕上显示正常
    try:
        from ctypes import windll
        windll.shcore.SetProcessDpiAwareness(1)
    except:
        pass

    root = tk.Tk()
    app = TranslatorApp(root)
    root.mainloop()